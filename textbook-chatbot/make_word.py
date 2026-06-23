#!/usr/bin/env python3
"""
Word文書生成ツール
使い方（JSON入力）:
  python3 ~/textbook-chatbot/make_word.py input.json
  python3 ~/textbook-chatbot/make_word.py input.json --out ~/Desktop/output.docx

JSON形式:
{
  "title": "文書タイトル",
  "subtitle": "サブタイトル（任意）",
  "sections": [
    {
      "heading": "見出し",
      "body": "本文テキスト",
      "bullets": ["箇条書き1", "箇条書き2"]
    }
  ],
  "table": {
    "headers": ["列1", "列2", "列3"],
    "rows": [["A","B","C"], ["D","E","F"]]
  },
  "references": ["参考文献1", "参考文献2"]
}
"""

import json, sys, argparse, re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = Path.home() / "Desktop"

# ── カラー定義 ─────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1d, 0x4e, 0xd8)   # ヘッダー色
BLUE_LIGHT = RGBColor(0xdb, 0xea, 0xff)   # 表ヘッダー背景
GRAY_TEXT  = RGBColor(0x44, 0x44, 0x44)   # 本文グレー


def set_cell_bg(cell, hex_color: str):
    """セルの背景色を設定"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int):
    """見出しを追加"""
    p = doc.add_heading(text, level=level)
    if level == 1:
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = BLUE_DARK
        run.font.size = Pt(16)
    elif level == 2:
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = BLUE_DARK
        run.font.size = Pt(13)
    return p


def add_body(doc: Document, text: str):
    """本文テキストを追加"""
    if not text.strip():
        return
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY_TEXT


def add_bullets(doc: Document, items: list):
    """箇条書きを追加"""
    for item in items:
        if not item.strip():
            continue
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY_TEXT


def add_table(doc: Document, headers: list, rows: list):
    """表を追加"""
    if not headers:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # ヘッダー行
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, "DBEAff")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = BLUE_DARK

    # データ行
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            if r_idx % 2 == 1:
                set_cell_bg(cell, "F5F7FF")
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)


def add_references(doc: Document, refs: list):
    """参考文献リストを追加"""
    if not refs:
        return
    doc.add_paragraph()
    add_heading(doc, "参考文献", 2)
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"・{ref}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def generate_word(data: dict, out_path: Path):
    """Word文書を生成"""
    doc = Document()

    # ── ページ設定 ─────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    # ── タイトル ───────────────────────────────────────────────
    title = data.get("title", "無題")
    p = doc.add_heading(title, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = BLUE_DARK
        run.font.size = Pt(20)

    if data.get("subtitle"):
        sp = doc.add_paragraph(data["subtitle"])
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sp.runs:
            run.font.size = Pt(12)
            run.font.color.rgb = GRAY_TEXT

    # 日付
    dp = doc.add_paragraph(datetime.now().strftime("%Y年%m月%d日"))
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in dp.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # ── セクション ─────────────────────────────────────────────
    for sec in data.get("sections", []):
        if sec.get("heading"):
            add_heading(doc, sec["heading"], 1)
        if sec.get("subheading"):
            add_heading(doc, sec["subheading"], 2)
        if sec.get("body"):
            add_body(doc, sec["body"])
        if sec.get("bullets"):
            add_bullets(doc, sec["bullets"])
        doc.add_paragraph()

    # ── 表 ────────────────────────────────────────────────────
    if data.get("table"):
        add_table(doc,
                  data["table"].get("headers", []),
                  data["table"].get("rows", []))
        doc.add_paragraph()

    # ── 参考文献 ───────────────────────────────────────────────
    add_references(doc, data.get("references", []))

    doc.save(str(out_path))
    return out_path


def main():
    parser = argparse.ArgumentParser(description="Word文書生成")
    parser.add_argument("input",  help="入力JSONファイル（または - でstdin）")
    parser.add_argument("--out", "-o", default="", help="出力ファイルパス")
    args = parser.parse_args()

    if args.input == "-":
        data = json.load(sys.stdin)
    else:
        with open(args.input, encoding="utf-8") as f:
            data = json.load(f)

    # 出力パスの決定
    if args.out:
        out_path = Path(args.out)
    else:
        ts   = datetime.now().strftime("%Y%m%d_%H%M")
        name = re.sub(r'[^\w぀-鿿]', '_', data.get("title", "output"))
        out_path = OUTPUT_DIR / f"{name}_{ts}.docx"

    out_path.parent.mkdir(parents=True, exist_ok=True)
    result = generate_word(data, out_path)
    print(f"✅ Word保存完了: {result}")


if __name__ == "__main__":
    main()
